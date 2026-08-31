from datetime import date
from decimal import Decimal
from io import BytesIO

import pytest
from docx import Document
from sqlmodel import select

from models import (
    Customer,
    CustomerType,
    DocumentArtifact,
    DocumentLegalEntity,
    DocumentNumberReservation,
    Order,
    OrderDocument,
    OrderProductLink,
    OrderProposal,
    OrderStatus,
    Product,
)
from models.tenancy import TenantScope
from modules.documents.application import (
    DocumentContextSelection,
    ManagedDocumentConflictError,
    ManagedDocumentGenerationError,
    ManagedDocumentService,
    NativeTemplatePlaceholderContract,
    NativeTemplateVersionService,
)
from modules.documents.domain import BusinessDocumentTerms, PaymentScheduleItem
from modules.documents.infrastructure.artifact_storage import (
    PrivateDocumentArtifactStorage,
)
from modules.documents.infrastructure.renderers import TableBlockSpec
from modules.documents.infrastructure.template_source_storage import (
    PrivateTemplateSourceStorage,
)
from services.private_attachment_storage_service import LocalPrivateAttachmentStorage


class FakePdfConverter:
    def health(self):
        raise AssertionError("health is not required for conversion")

    def convert_docx(self, content: bytes, *, filename: str = "document.docx") -> bytes:
        assert content.startswith(b"PK")
        assert filename.endswith(".docx")
        return b"%PDF-1.4\nmanaged-document-test\n%%EOF"


class FailingPdfConverter(FakePdfConverter):
    def convert_docx(self, content: bytes, *, filename: str = "document.docx") -> bytes:
        raise RuntimeError("converter unavailable")


def _template_docx() -> bytes:
    document = Document()
    document.add_paragraph("Договор № {{ document.official_full_number }}")
    document.add_paragraph("Покупатель: {{ customer.full_name }}")
    table = document.add_table(rows=2, cols=4)
    for index, value in enumerate(("№", "Наименование", "Кол-во", "Сумма")):
        table.rows[0].cells[index].text = value
    for index, value in enumerate(
        (
            "{{ lines }}{{ line.number }}",
            "{{ line.title }}",
            "{{ line.quantity }}",
            "{{ line.amount }}",
        )
    ):
        table.rows[1].cells[index].text = value
    output = BytesIO()
    document.save(output)
    return output.getvalue()


async def _seed(db, tmp_path):
    customer = Customer(
        tenant_id=1,
        name="ООО Клиент",
        full_legal_name="Общество с ограниченной ответственностью Клиент",
        phone="+375290000001",
        type=CustomerType.company,
    )
    product = Product(
        title="Кондиционер",
        slug="managed-document-lifecycle-conditioner",
        price=1_500,
    )
    issuer = DocumentLegalEntity(
        tenant_id=1,
        slug="managed-documents",
        display_name="ООО Продавец",
        legal_name="Общество с ограниченной ответственностью Продавец",
        unp="390000001",
        is_default=True,
        requisites={},
    )
    db.add_all([customer, product, issuer])
    await db.flush()
    order = Order(
        tenant_id=1,
        storefront_id=1,
        customer_id=customer.id,
        status=OrderStatus.NEGOTIATION,
        title="Поставка кондиционера",
        total_amount=1_500,
        total_cost=1_000,
        margin=500,
    )
    db.add(order)
    await db.flush()
    proposal = OrderProposal(
        order_id=order.id,
        name="Основное предложение",
        is_selected=True,
    )
    db.add(proposal)
    await db.flush()
    db.add(
        OrderProductLink(
            order_id=order.id,
            proposal_id=proposal.id,
            product_id=product.id,
            quantity=1,
            price=1_500,
            cost=1_000,
        )
    )
    await db.commit()

    scope = TenantScope(tenant_id=1, storefront_id=1, is_system=True)
    private = LocalPrivateAttachmentStorage(tmp_path / "private-documents")
    template_storage = PrivateTemplateSourceStorage(private)
    artifact_storage = PrivateDocumentArtifactStorage(private)
    template = await NativeTemplateVersionService.create_template(
        db,
        tenant_scope=scope,
        legal_entity_id=issuer.id,
        name="Договор поставки",
        doc_type="contract",
    )
    version = await NativeTemplateVersionService.upload_native_docx_version(
        db,
        tenant_scope=scope,
        legal_entity_id=issuer.id,
        template_id=template.id,
        filename="Договор.docx",
        content=_template_docx(),
        placeholder_contract=NativeTemplatePlaceholderContract.create(
            field_catalog={
                "document.official_full_number",
                "customer.full_name",
            },
            table_blocks=(
                TableBlockSpec(
                    name="lines",
                    row_fields=frozenset(
                        {
                            "line.number",
                            "line.title",
                            "line.quantity",
                            "line.amount",
                        }
                    ),
                ),
            ),
        ),
        storage=template_storage,
    )
    await NativeTemplateVersionService.activate_version(
        db,
        tenant_scope=scope,
        legal_entity_id=issuer.id,
        template_id=template.id,
        version_id=version.id,
    )
    return scope, order, issuer, template_storage, artifact_storage


async def _draft(
    db, *, scope, order, issuer, issue_date=date(2026, 8, 26), replaces=None
):
    return await ManagedDocumentService.create_draft(
        db,
        tenant_scope=scope,
        selection=DocumentContextSelection(
            order_id=order.id,
            legal_entity_id=issuer.id,
            document_type="contract",
            issue_date=issue_date,
            business_terms=BusinessDocumentTerms(
                contract_scenario="services",
                payment_schedule=(
                    PaymentScheduleItem(Decimal("100"), "before_work"),
                ),
            ),
        ),
        replaces_document_id=replaces,
    )


@pytest.mark.asyncio
async def test_issue_is_idempotent_and_persists_frozen_docx_pdf_and_number(
    db, tmp_path
):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    draft = await _draft(db, scope=scope, order=order, issuer=issuer)

    assert draft.status == "draft"
    assert draft.official_number is None
    result = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )

    assert result.document.status == "issued"
    assert result.document.official_series == "Д-"
    assert result.document.official_period_key == "2026"
    assert result.document.official_number == "001"
    assert (
        result.document.render_snapshot["values"]["document.official_full_number"]
        == "Д-001"
    )
    assert {artifact.kind for artifact in result.artifacts} == {"rendered_docx", "pdf"}
    assert all(artifact.is_authoritative for artifact in result.artifacts)

    repeated = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FailingPdfConverter(),
    )
    assert repeated.document.official_number == "001"
    assert len(repeated.artifacts) == 2
    assert (
        len((await db.execute(select(DocumentNumberReservation))).scalars().all()) == 1
    )


@pytest.mark.asyncio
async def test_failed_conversion_keeps_number_for_safe_retry_and_void_never_reuses_it(
    db, tmp_path
):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    failed = await _draft(db, scope=scope, order=order, issuer=issuer)

    with pytest.raises(ManagedDocumentGenerationError, match="сохранил номер"):
        await ManagedDocumentService.issue(
            db,
            tenant_scope=scope,
            document_id=failed.id,
            template_storage=template_storage,
            artifact_storage=artifact_storage,
            pdf_converter=FailingPdfConverter(),
        )
    await db.refresh(failed)
    assert failed.status == "draft"
    assert failed.official_number == "001"
    assert (await db.execute(select(DocumentArtifact))).scalars().all() == []

    retried = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=failed.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )
    assert retried.document.official_number == "001"
    voided = await ManagedDocumentService.void(
        db,
        tenant_scope=scope,
        document_id=failed.id,
        reason="Исправлена существенная ошибка",
    )
    assert voided.status == "void"

    following = await _draft(db, scope=scope, order=order, issuer=issuer)
    following_result = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=following.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )
    assert following_result.document.official_number == "002"
    reservations = (
        (
            await db.execute(
                select(DocumentNumberReservation).order_by(
                    DocumentNumberReservation.number_value
                )
            )
        )
        .scalars()
        .all()
    )
    assert [(row.number_value, row.status) for row in reservations] == [
        (1, "void"),
        (2, "assigned"),
    ]


@pytest.mark.asyncio
async def test_calendar_year_scope_and_replacement_chain(db, tmp_path):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    original = await _draft(db, scope=scope, order=order, issuer=issuer)
    original_result = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=original.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )
    replacement = await _draft(
        db,
        scope=scope,
        order=order,
        issuer=issuer,
        replaces=original.id,
    )
    replacement_result = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=replacement.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )
    await db.refresh(original)
    assert original_result.document.id == original.id
    assert original.status == "replaced"
    assert replacement_result.document.replaces_document_id == original.id
    assert replacement_result.document.official_number == "002"

    next_year = await _draft(
        db,
        scope=scope,
        order=order,
        issuer=issuer,
        issue_date=date(2027, 1, 3),
    )
    next_year_result = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=next_year.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )
    assert next_year_result.document.official_period_key == "2027"
    assert next_year_result.document.official_number == "001"
    assert len((await db.execute(select(OrderDocument))).scalars().all()) == 3


@pytest.mark.asyncio
async def test_only_one_active_replacement_can_exist_for_a_document(db, tmp_path):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    original = await _draft(db, scope=scope, order=order, issuer=issuer)
    await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=original.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )

    first_replacement = await _draft(
        db,
        scope=scope,
        order=order,
        issuer=issuer,
        replaces=original.id,
    )
    with pytest.raises(ManagedDocumentConflictError, match="уже готовится"):
        await _draft(
            db,
            scope=scope,
            order=order,
            issuer=issuer,
            replaces=original.id,
        )

    issued = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=first_replacement.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )
    assert issued.document.replaces_document_id == original.id
