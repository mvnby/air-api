from datetime import datetime, timezone
from io import BytesIO

import pytest
from docx import Document

from modules.documents.application.managed_document_external_edits import (
    ManagedDocumentExternalEditSessionService,
)
from modules.documents.application.errors import ManagedDocumentConflictError
from modules.documents.application.external_edit_sessions import (
    ExternalEditSessionConflictError,
)
from modules.documents.application.lifecycle_service import ManagedDocumentService
from modules.documents.application.editable_draft_issue import (
    verify_document_external_edit_before_issue,
)
from modules.documents.infrastructure.external_edit_provider import (
    DOCX_CONTENT_TYPE,
    DownloadedExternalEditFile,
    ExternalEditFileMetadata,
)
from tests.integration.test_managed_document_lifecycle import (
    FakePdfConverter,
    _draft,
    _seed,
)


class FakeGoogleEditor:
    provider_name = "google_drive"
    connection_id = "connection-1"

    def __init__(self) -> None:
        self.content = b""
        self.revision = "revision-1"
        self.filename = "draft.docx"
        self.edit_session_id = "pending"

    def metadata(self) -> ExternalEditFileMetadata:
        return ExternalEditFileMetadata(
            file_id="google-file-1",
            edit_session_id=self.edit_session_id,
            edit_url="https://docs.google.com/document/d/google-file-1/edit",
            filename=self.filename,
            mime_type=DOCX_CONTENT_TYPE,
            revision=self.revision,
            modified_at=datetime.now(timezone.utc),
        )

    async def ensure_docx(self, *, edit_session_id: str, filename: str, content: bytes):
        assert edit_session_id
        self.edit_session_id = edit_session_id
        if not self.content:
            self.content = content
            self.filename = filename
        return self.metadata()

    async def get_metadata(self, file_id: str):
        assert file_id == "google-file-1"
        return self.metadata()

    async def download_docx(self, file_id: str):
        assert file_id == "google-file-1"
        return DownloadedExternalEditFile(self.metadata(), self.content)

    def add_manual_paragraph(self, text: str) -> None:
        document = Document(BytesIO(self.content))
        document.add_paragraph(text)
        output = BytesIO()
        document.save(output)
        self.content = output.getvalue()
        self.revision = "revision-2"


def _document_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        paragraphs.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(paragraphs)


@pytest.mark.asyncio
async def test_google_edited_draft_keeps_manual_changes_when_number_is_issued(
    db, tmp_path
):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    draft = await _draft(db, scope=scope, order=order, issuer=issuer)
    provider = FakeGoogleEditor()

    edit_session = await ManagedDocumentExternalEditSessionService.ensure_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        provider=provider,
    )
    editable_text = _document_text(provider.content)
    assert "Общество с ограниченной ответственностью Клиент" in editable_text
    assert "{{ document.official_full_number }}" in editable_text
    original_checksum = edit_session.base_checksum_sha256

    provider.add_manual_paragraph("Согласовано вручную в Google Docs")
    changed = await ManagedDocumentExternalEditSessionService.get_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        provider=provider,
    )
    assert changed.status == "changed"

    synced = await ManagedDocumentExternalEditSessionService.sync(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        expected_base_checksum_sha256=original_checksum,
        expected_remote_revision=changed.remote_revision,
        idempotency_key="sync-google-draft-1",
        artifact_storage=artifact_storage,
        provider=provider,
    )
    assert synced.status == "ready"
    assert synced.base_checksum_sha256 != original_checksum

    verified_revision = await verify_document_external_edit_before_issue(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        provider=provider,
    )
    issued = await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
        verified_remote_revision=verified_revision,
    )
    rendered = next(item for item in issued.artifacts if item.kind == "rendered_docx")
    rendered_content = await artifact_storage.read(
        ManagedDocumentService.stored_artifact(rendered)
    )
    final_text = _document_text(rendered_content)
    assert "Согласовано вручную в Google Docs" in final_text
    assert "Д-2026-001" in final_text
    assert "{{ document.official_full_number }}" not in final_text


@pytest.mark.asyncio
async def test_issue_preflight_detects_remote_change_without_prior_status_refresh(
    db, tmp_path
):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    draft = await _draft(db, scope=scope, order=order, issuer=issuer)
    provider = FakeGoogleEditor()
    edit_session = await ManagedDocumentExternalEditSessionService.ensure_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        provider=provider,
    )
    assert edit_session.status == "ready"

    provider.add_manual_paragraph("Правка прямо перед выпуском")

    with pytest.raises(ValueError, match="несинхронизированные"):
        await verify_document_external_edit_before_issue(
            db,
            tenant_scope=scope,
            document_id=draft.id,
            provider=provider,
        )

    await db.refresh(edit_session)
    assert edit_session.status == "changed"
    assert draft.official_number is None


@pytest.mark.asyncio
async def test_issued_document_cannot_start_google_edit_session(db, tmp_path):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    draft = await _draft(db, scope=scope, order=order, issuer=issuer)
    await ManagedDocumentService.issue(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        pdf_converter=FakePdfConverter(),
    )

    with pytest.raises(ValueError, match="только черновик"):
        await ManagedDocumentExternalEditSessionService.ensure_session(
            db,
            tenant_scope=scope,
            document_id=draft.id,
            template_storage=template_storage,
            artifact_storage=artifact_storage,
            provider=FakeGoogleEditor(),
        )


@pytest.mark.asyncio
async def test_issue_blocks_unsynced_google_changes_before_reserving_number(
    db, tmp_path
):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    draft = await _draft(db, scope=scope, order=order, issuer=issuer)
    provider = FakeGoogleEditor()
    await ManagedDocumentExternalEditSessionService.ensure_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        provider=provider,
    )
    provider.add_manual_paragraph("Ещё не синхронизировано")
    changed = await ManagedDocumentExternalEditSessionService.get_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        provider=provider,
    )
    assert changed.status == "changed"

    with pytest.raises(ManagedDocumentConflictError, match="служебные поля"):
        await ManagedDocumentService.issue(
            db,
            tenant_scope=scope,
            document_id=draft.id,
            template_storage=template_storage,
            artifact_storage=artifact_storage,
            pdf_converter=FakePdfConverter(),
        )

    await db.refresh(draft)
    assert draft.official_number is None


@pytest.mark.asyncio
async def test_google_reconnect_creates_new_auditable_document_session(db, tmp_path):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    draft = await _draft(db, scope=scope, order=order, issuer=issuer)
    original_provider = FakeGoogleEditor()
    original = await ManagedDocumentExternalEditSessionService.ensure_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        provider=original_provider,
    )

    reconnected_provider = FakeGoogleEditor()
    reconnected_provider.connection_id = "connection-2"
    current = await ManagedDocumentExternalEditSessionService.ensure_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        provider=reconnected_provider,
    )

    assert current.id != original.id
    assert current.provider_connection_id == "connection-2"
    preserved = await db.get(type(original), original.id)
    assert preserved is not None
    assert preserved.provider_connection_id == "connection-1"


@pytest.mark.asyncio
async def test_document_sync_idempotency_key_is_bound_to_payload(db, tmp_path):
    scope, order, issuer, template_storage, artifact_storage = await _seed(db, tmp_path)
    draft = await _draft(db, scope=scope, order=order, issuer=issuer)
    provider = FakeGoogleEditor()
    edit_session = await ManagedDocumentExternalEditSessionService.ensure_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        template_storage=template_storage,
        artifact_storage=artifact_storage,
        provider=provider,
    )
    provider.add_manual_paragraph("Первая редакция")
    changed = await ManagedDocumentExternalEditSessionService.get_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        provider=provider,
    )
    synced = await ManagedDocumentExternalEditSessionService.sync(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        expected_base_checksum_sha256=edit_session.base_checksum_sha256,
        expected_remote_revision=changed.remote_revision,
        idempotency_key="one-logical-sync",
        artifact_storage=artifact_storage,
        provider=provider,
    )
    provider.add_manual_paragraph("Вторая редакция")
    changed_again = await ManagedDocumentExternalEditSessionService.get_session(
        db,
        tenant_scope=scope,
        document_id=draft.id,
        provider=provider,
    )

    with pytest.raises(ExternalEditSessionConflictError, match="другой версии"):
        await ManagedDocumentExternalEditSessionService.sync(
            db,
            tenant_scope=scope,
            document_id=draft.id,
            expected_base_checksum_sha256=synced.base_checksum_sha256,
            expected_remote_revision=changed_again.remote_revision,
            idempotency_key="one-logical-sync",
            artifact_storage=artifact_storage,
            provider=provider,
        )
