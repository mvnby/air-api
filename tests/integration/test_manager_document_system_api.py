import json
from datetime import timedelta
from io import BytesIO

import pytest
from docx import Document
from httpx import AsyncClient

from core.config import settings
from core.security import create_access_token
from models import (
    Customer,
    DocumentArtifact,
    Order,
    OrderStatus,
    StaffUser,
    Storefront,
    Tenant,
    TenantMembership,
)
from services.private_attachment_storage_service import LocalPrivateAttachmentStorage


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
BASE = "/api/manager/document-system"


class _FakePdfConverter:
    def health(self):  # pragma: no cover - the issue endpoint does not call it
        raise AssertionError("health is not used while issuing a document")

    def convert_docx(self, content: bytes, *, filename: str = "document.docx") -> bytes:
        assert content.startswith(b"PK")
        assert filename.endswith(".docx")
        return b"%PDF-1.4\napi-document-system-test\n%%EOF"


@pytest.fixture(autouse=True)
def _local_document_storage(monkeypatch, tmp_path):
    """Keep document-system API tests independent of configured S3/R2 clients."""
    import importlib

    document_router = importlib.import_module("modules.documents.api.router")
    private = LocalPrivateAttachmentStorage(tmp_path / "private-document-system-api")
    monkeypatch.setattr(
        document_router, "get_private_attachment_storage", lambda provider=None: private
    )
    return private


async def _legacy_owner_headers(client: AsyncClient) -> dict[str, str]:
    response = await client.post(
        "/login/access-token",
        data={"username": settings.ADMIN_USERNAME, "password": settings.ADMIN_PASSWORD},
    )
    assert response.status_code == 200
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


async def _create_tenant_owner(
    db,
    *,
    slug: str,
    username: str,
    role: str = "owner",
) -> tuple[Tenant, Storefront, StaffUser]:
    tenant = Tenant(
        slug=slug, display_name=slug.title(), status="active", is_system=False
    )
    db.add(tenant)
    await db.flush()
    storefront = Storefront(
        tenant_id=int(tenant.id),
        slug="main",
        display_name=f"{slug.title()} Main",
        status="active",
        is_default=True,
    )
    user = StaffUser(
        display_name=username,
        status="active",
        roles=[role],
        primary_role=role,
        username=username,
    )
    db.add_all([storefront, user])
    await db.flush()
    db.add(
        TenantMembership(
            tenant_id=int(tenant.id),
            staff_user_id=int(user.id),
            role=role,
            status="active",
        )
    )
    await db.commit()
    return tenant, storefront, user


def _staff_headers(user: StaffUser) -> dict[str, str]:
    token = create_access_token(
        {
            "sub": user.username,
            "staff_user_id": user.id,
            "auth_version": user.auth_version,
            "auth_source": "document-system-api-test",
        },
        expires_delta=timedelta(minutes=10),
    )
    return {"Authorization": f"Bearer {token}"}


def _template_docx() -> bytes:
    document = Document()
    document.add_paragraph("Договор № {{ document.official_full_number }}")
    document.add_paragraph("Заказчик: {{ customer.full_name }}")
    document.add_paragraph("{{#if seller.is_individual_entrepreneur}}")
    document.add_paragraph("Продавец действует как ИП")
    document.add_paragraph("{{/if seller.is_individual_entrepreneur}}")
    table = document.add_table(rows=2, cols=3)
    for index, value in enumerate(("Этап", "%", "Сумма")):
        table.rows[0].cells[index].text = value
    for index, value in enumerate(
        (
            "{{ payment_schedule }}{{ payment.number }}",
            "{{ payment.share_percent }}",
            "{{ payment.amount }}",
        )
    ):
        table.rows[1].cells[index].text = value
    output = BytesIO()
    document.save(output)
    return output.getvalue()


async def _create_issuer(
    client: AsyncClient, headers: dict[str, str], *, name: str
) -> int:
    response = await client.post(
        f"{BASE}/legal-entities",
        headers=headers,
        json={"display_name": name, "slug": name.lower().replace(" ", "-")},
    )
    assert response.status_code == 200, response.text
    return int(response.json()["id"])


async def _create_native_contract_template(
    client: AsyncClient,
    headers: dict[str, str],
    *,
    legal_entity_id: int,
) -> tuple[int, int]:
    created = await client.post(
        f"{BASE}/templates",
        headers=headers,
        json={
            "legal_entity_id": legal_entity_id,
            "name": "Договор поставки",
            "doc_type": "contract",
        },
    )
    assert created.status_code == 200, created.text
    template_id = int(created.json()["id"])
    uploaded = await client.post(
        f"{BASE}/templates/{template_id}/versions",
        headers=headers,
        data={"legal_entity_id": str(legal_entity_id)},
        files={"file": ("Договор.docx", _template_docx(), DOCX_MIME)},
    )
    assert uploaded.status_code == 200, uploaded.text
    version_id = int(uploaded.json()["id"])
    activated = await client.post(
        f"{BASE}/templates/{template_id}/versions/{version_id}/activate",
        headers=headers,
        params={"legal_entity_id": legal_entity_id},
    )
    assert activated.status_code == 200, activated.text
    assert activated.json()["status"] == "active"
    return template_id, version_id


async def _seed_order(db) -> Order:
    customer = Customer(
        tenant_id=1,
        name="ООО Клиент API",
        full_legal_name="Общество с ограниченной ответственностью Клиент API",
        phone="+375290000101",
    )
    db.add(customer)
    await db.flush()
    order = Order(
        tenant_id=1,
        storefront_id=1,
        customer_id=int(customer.id),
        status=OrderStatus.NEGOTIATION,
        title="Поставка для API",
    )
    db.add(order)
    await db.commit()
    return order


@pytest.mark.asyncio
async def test_document_system_native_template_flow_discovers_catalog_and_activates(
    async_client: AsyncClient,
):
    headers = await _legacy_owner_headers(async_client)

    catalog = await async_client.get(
        f"{BASE}/placeholder-catalog",
        headers=headers,
        params={"doc_type": "contract"},
    )
    assert catalog.status_code == 200
    assert "document.official_full_number" in {
        item["name"] for item in catalog.json()["fields"]
    }
    assert "seller.entity_type_label" in {
        item["name"] for item in catalog.json()["fields"]
    }
    assert "document.issue_city" in {item["name"] for item in catalog.json()["fields"]}
    assert "seller.is_individual_entrepreneur" in {
        item["name"] for item in catalog.json()["conditions"]
    }
    seller_ip_condition = next(
        item
        for item in catalog.json()["conditions"]
        if item["name"] == "seller.is_individual_entrepreneur"
    )
    assert seller_ip_condition["start_syntax"] == (
        "{{#if seller.is_individual_entrepreneur}}"
    )
    assert seller_ip_condition["end_syntax"] == (
        "{{/if seller.is_individual_entrepreneur}}"
    )
    assert [item["anchor_syntax"] for item in catalog.json()["tables"]] == [
        "{{ lines }}",
        "{{ payment_schedule }}",
    ]
    waybill_catalog = await async_client.get(
        f"{BASE}/placeholder-catalog",
        headers=headers,
        params={"doc_type": "tn2"},
    )
    assert waybill_catalog.status_code == 200
    assert "transport.car_number" in {
        item["name"] for item in waybill_catalog.json()["fields"]
    }
    assert "line.country" in {
        item["name"] for item in waybill_catalog.json()["tables"][0]["row_fields"]
    }

    issuer_id = await _create_issuer(async_client, headers, name="ООО API Продавец")
    template_id, version_id = await _create_native_contract_template(
        async_client,
        headers,
        legal_entity_id=issuer_id,
    )
    versions = await async_client.get(
        f"{BASE}/templates/{template_id}/versions",
        headers=headers,
        params={"legal_entity_id": issuer_id},
    )
    assert versions.status_code == 200, versions.text
    assert len(versions.json()["items"]) == 1
    version = versions.json()["items"][0]
    assert version["id"] == version_id
    assert version["version"] == 1
    assert version["status"] == "active"
    assert version["renderer"] == "docx"
    assert version["placeholder_schema"] == {
        "fields": ["customer.full_name", "document.official_full_number"],
        "conditions": ["seller.is_individual_entrepreneur"],
        "tables": [
            {
                "name": "payment_schedule",
                "row_fields": [
                    "payment.amount",
                    "payment.due_day_kind",
                    "payment.due_days",
                    "payment.due_event",
                    "payment.note",
                    "payment.number",
                    "payment.share_percent",
                ],
            }
        ],
    }
    source = await async_client.get(
        f"{BASE}/templates/{template_id}/versions/{version_id}/source",
        headers=headers,
        params={"legal_entity_id": issuer_id},
    )
    assert source.status_code == 200, source.text
    assert source.content.startswith(b"PK")
    assert source.headers["cache-control"] == "private, no-store"
    assert "filename*=UTF-8''" in source.headers["content-disposition"]

    manual_schema = {
        "fields": ["customer.full_name", "document.official_full_number"],
        "conditions": ["seller.is_individual_entrepreneur"],
        "tables": [
            {
                "name": "payment_schedule",
                "row_fields": [
                    "payment.amount",
                    "payment.number",
                    "payment.share_percent",
                ],
            }
        ],
    }
    manually_uploaded = await async_client.post(
        f"{BASE}/templates/{template_id}/versions",
        headers=headers,
        data={
            "legal_entity_id": str(issuer_id),
            "placeholder_schema": json.dumps(manual_schema),
        },
        files={"file": ("Договор-2.docx", _template_docx(), DOCX_MIME)},
    )
    assert manually_uploaded.status_code == 200, manually_uploaded.text
    assert manually_uploaded.json()["placeholder_schema"] == manual_schema


@pytest.mark.asyncio
async def test_document_system_persists_seller_entity_type(async_client: AsyncClient):
    headers = await _legacy_owner_headers(async_client)

    created = await async_client.post(
        f"{BASE}/legal-entities",
        headers=headers,
        json={
            "display_name": "ИП API Продавец",
            "entity_type": "individual_entrepreneur",
        },
    )
    assert created.status_code == 200, created.text
    assert created.json()["entity_type"] == "individual_entrepreneur"

    updated = await async_client.patch(
        f"{BASE}/legal-entities/{created.json()['id']}",
        headers=headers,
        json={"entity_type": "organization"},
    )
    assert updated.status_code == 200, updated.text
    assert updated.json()["entity_type"] == "organization"


@pytest.mark.asyncio
async def test_document_system_settings_mutations_require_owner(
    async_client: AsyncClient,
    db,
):
    _tenant, _storefront, manager = await _create_tenant_owner(
        db,
        slug="document-manager",
        username="document-manager",
        role="manager",
    )
    headers = _staff_headers(manager)

    read_catalog = await async_client.get(
        f"{BASE}/placeholder-catalog",
        headers=headers,
        params={"doc_type": "invoice"},
    )
    write_template = await async_client.post(
        f"{BASE}/templates",
        headers=headers,
        json={"legal_entity_id": 999, "name": "Недоступный", "doc_type": "contract"},
    )
    write_issuer = await async_client.post(
        f"{BASE}/legal-entities",
        headers=headers,
        json={"display_name": "Недоступное юридическое лицо"},
    )

    assert read_catalog.status_code == 200
    assert write_template.status_code == 403
    assert write_issuer.status_code == 403


@pytest.mark.asyncio
async def test_document_system_draft_issue_is_idempotent_and_artifacts_are_tenant_isolated(
    async_client: AsyncClient,
    db,
    tmp_path,
    monkeypatch,
):
    import importlib

    document_router = importlib.import_module("modules.documents.api.router")
    monkeypatch.setattr(document_router, "_pdf_converter", lambda: _FakePdfConverter())

    headers = await _legacy_owner_headers(async_client)
    issuer_id = await _create_issuer(async_client, headers, name="ООО Выпуск API")
    template_id, _version_id = await _create_native_contract_template(
        async_client,
        headers,
        legal_entity_id=issuer_id,
    )
    order = await _seed_order(db)
    draft = await async_client.post(
        f"{BASE}/orders/{order.id}/documents/drafts",
        headers=headers,
        json={
            "legal_entity_id": issuer_id,
            "document_type": "contract",
            "issue_date": "2026-08-26",
            "issue_city": "Минск",
            "template_id": template_id,
            "business_terms": {
                "contract_scenario": "services",
                "payment_schedule": [
                    {"share_percent": 100, "due_event": "before_work"}
                ],
            },
        },
    )
    assert draft.status_code == 200, draft.text
    document_id = int(draft.json()["id"])
    assert draft.json()["status"] == "draft"
    assert draft.json()["official_number"] is None
    assert draft.json()["issue_city"] == "Минск"

    issued = await async_client.post(
        f"{BASE}/documents/{document_id}/issue", headers=headers
    )
    repeated = await async_client.post(
        f"{BASE}/documents/{document_id}/issue", headers=headers
    )
    assert issued.status_code == repeated.status_code == 200
    assert (
        issued.json()["official_full_number"]
        == repeated.json()["official_full_number"]
        == "Д-001"
    )
    assert {artifact["kind"] for artifact in issued.json()["artifacts"]} == {
        "pdf",
        "rendered_docx",
    }
    artifact_id = issued.json()["artifacts"][0]["id"]
    own_download = await async_client.get(
        f"{BASE}/artifacts/{artifact_id}/download",
        headers=headers,
    )
    assert own_download.status_code == 200
    assert own_download.headers["cache-control"] == "private, no-store"

    artifact = await db.get(DocumentArtifact, artifact_id)
    assert artifact is not None
    artifact.checksum_sha256 = "0" * 64
    db.add(artifact)
    await db.commit()
    corrupted = await async_client.get(
        f"{BASE}/artifacts/{artifact_id}/download",
        headers=headers,
    )
    assert corrupted.status_code == 409
    assert corrupted.json()["detail"]["error_code"] == (
        "document_artifact_integrity_failed"
    )

    _tenant_b, _storefront_b, owner_b = await _create_tenant_owner(
        db,
        slug="document-artifact-other",
        username="document-artifact-other-owner",
    )
    denied = await async_client.get(
        f"{BASE}/artifacts/{artifact_id}/download",
        headers=_staff_headers(owner_b),
    )
    assert denied.status_code == 404


@pytest.mark.asyncio
async def test_document_system_templates_are_hidden_from_another_tenant(
    async_client: AsyncClient,
    db,
):
    headers_a = await _legacy_owner_headers(async_client)
    issuer_a = await _create_issuer(async_client, headers_a, name="ООО Tenant A")
    template_id, _version_id = await _create_native_contract_template(
        async_client,
        headers_a,
        legal_entity_id=issuer_a,
    )
    _tenant_b, _storefront_b, owner_b = await _create_tenant_owner(
        db,
        slug="document-template-other",
        username="document-template-other-owner",
    )
    headers_b = _staff_headers(owner_b)
    issuer_b = await _create_issuer(async_client, headers_b, name="ООО Tenant B")

    foreign_list = await async_client.get(
        f"{BASE}/templates/{template_id}/versions",
        headers=headers_b,
        params={"legal_entity_id": issuer_b},
    )
    foreign_template_list = await async_client.get(
        f"{BASE}/templates",
        headers=headers_b,
        params={"legal_entity_id": issuer_a},
    )

    assert foreign_list.status_code == 404
    assert foreign_template_list.status_code == 404
