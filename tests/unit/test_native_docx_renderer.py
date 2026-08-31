from io import BytesIO

import pytest
from docx import Document
from docx.oxml.ns import qn

from modules.documents.infrastructure.renderers import (
    ContextFieldError,
    DocumentTemplateVersion,
    GotenbergPdfConverter,
    NativeDocxRenderer,
    RenderContext,
    TableBlockSpec,
    TemplateValidationError,
    UnavailablePdfConverter,
)
from modules.documents.infrastructure.renderers.pdf import (
    PdfConversionError,
    PdfConversionUnavailableError,
)


def _template_bytes(*, unknown_placeholder: str | None = None) -> bytes:
    document = Document()
    heading = document.add_paragraph("Документ № ")
    heading.add_run("{{ document.")
    heading.add_run("official_number }}")
    document.add_paragraph("Клиент: {{ customer.full_name }}")
    if unknown_placeholder:
        document.add_paragraph(f"Скрытое: {{{{ {unknown_placeholder} }}}}")
    table = document.add_table(rows=1, cols=4)
    row = table.rows[0].cells
    row[0].text = "{{ lines }}{{ line.number }}"
    row[1].text = "{{ line.title }}"
    row[2].text = "{{ line.quantity }}"
    row[3].text = "{{ line.amount }}"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _template(source: bytes) -> DocumentTemplateVersion:
    return DocumentTemplateVersion(
        template_key="invoice",
        version=3,
        source=source,
        field_catalog=frozenset({"document.official_number", "customer.full_name"}),
        table_blocks=(
            TableBlockSpec(
                name="lines",
                row_fields=frozenset(
                    {"line.number", "line.title", "line.quantity", "line.amount"}
                ),
            ),
        ),
    )


def _conditional_template(
    source: bytes,
    *,
    conditions: frozenset[str] = frozenset({"seller.is_individual_entrepreneur"}),
) -> DocumentTemplateVersion:
    return DocumentTemplateVersion(
        template_key="contract",
        version=1,
        source=source,
        field_catalog=frozenset({"seller.city"}),
        condition_catalog=conditions,
    )


def _save(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_native_renderer_replaces_split_placeholders_and_repeats_table_rows():
    rendered = NativeDocxRenderer().render(
        _template(_template_bytes()),
        RenderContext(
            values={
                "document.official_number": "С-2026-001",
                "customer.full_name": "ООО Тест",
            },
            table_rows={
                "lines": (
                    {
                        "line.number": "1",
                        "line.title": "Кондиционер",
                        "line.quantity": "1",
                        "line.amount": "2500.00",
                    },
                    {
                        "line.number": "2",
                        "line.title": "Монтаж",
                        "line.quantity": "1",
                        "line.amount": "500.00",
                    },
                )
            },
        ),
    )

    result = Document(BytesIO(rendered.content))
    assert [paragraph.text for paragraph in result.paragraphs] == [
        "Документ № С-2026-001",
        "Клиент: ООО Тест",
    ]
    assert [[cell.text for cell in row.cells] for row in result.tables[0].rows] == [
        ["1", "Кондиционер", "1", "2500.00"],
        ["2", "Монтаж", "1", "500.00"],
    ]
    assert rendered.template_key == "invoice"
    assert rendered.template_version == 3


def test_renderer_removes_only_empty_trailing_paragraph_after_final_table():
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Подпись"
    document.add_paragraph()
    template = DocumentTemplateVersion(
        template_key="contract",
        version=1,
        source=_save(document),
        field_catalog=frozenset(),
    )

    rendered = NativeDocxRenderer().render(template, RenderContext(values={}))

    result = Document(BytesIO(rendered.content))
    body_content_tags = [
        child.tag
        for child in result._element.body.iterchildren()
        if child.tag != result._element.body.sectPr.tag
    ]
    assert body_content_tags == [result.tables[0]._element.tag]


def test_renderer_keeps_trailing_paragraph_with_a_page_break_after_final_table():
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Подпись"
    document.add_paragraph().add_run().add_break()
    template = DocumentTemplateVersion(
        template_key="contract",
        version=1,
        source=_save(document),
        field_catalog=frozenset(),
    )

    rendered = NativeDocxRenderer().render(template, RenderContext(values={}))

    result = Document(BytesIO(rendered.content))
    assert len(result.paragraphs) == 1


def test_renderer_preserves_multiline_clause_as_word_line_breaks():
    document = Document()
    document.add_paragraph("Условия: {{ contract.additional_conditions }}")
    template = DocumentTemplateVersion(
        template_key="contract",
        version=1,
        source=_save(document),
        field_catalog=frozenset({"contract.additional_conditions"}),
    )

    rendered = NativeDocxRenderer().render(
        template,
        RenderContext(
            values={"contract.additional_conditions": "Первое условие\nВторое условие"}
        ),
    )

    result = Document(BytesIO(rendered.content))
    assert result.paragraphs[0].text == "Условия: Первое условие\nВторое условие"
    assert len(list(result.paragraphs[0]._p.iter(qn("w:br")))) == 1


def test_renderer_discovers_template_placeholders_without_expanding_catalogue():
    discovered = NativeDocxRenderer().discover_placeholders(_template_bytes())

    assert discovered == {
        "document.official_number",
        "customer.full_name",
        "lines",
        "line.number",
        "line.title",
        "line.quantity",
        "line.amount",
    }


def test_renderer_refuses_unapproved_context_fields_before_any_document_is_emitted():
    context = RenderContext(
        values={
            "document.official_number": "С-2026-001",
            "customer.full_name": "ООО Тест",
            "internal.audit_secret": "classified",
        },
        table_rows={"lines": ()},
    )

    with pytest.raises(ContextFieldError, match="internal.audit_secret"):
        NativeDocxRenderer().render(_template(_template_bytes()), context)


def test_render_context_is_frozen_after_the_document_snapshot_is_built():
    context = RenderContext(
        values={"document.official_number": "С-2026-001"},
        table_rows={"lines": ({"line.number": "1"},)},
    )

    with pytest.raises(TypeError):
        context.values["document.official_number"] = "С-2026-002"
    with pytest.raises(TypeError):
        context.table_rows["lines"][0]["line.number"] = "2"


def test_renderer_refuses_unknown_template_placeholder_without_leaking_context_data():
    context = RenderContext(
        values={
            "document.official_number": "С-2026-001",
            "customer.full_name": "ООО Тест",
        },
        table_rows={"lines": ()},
    )

    with pytest.raises(TemplateValidationError, match="internal.secret"):
        NativeDocxRenderer().render(
            _template(_template_bytes(unknown_placeholder="internal.secret")), context
        )


@pytest.mark.parametrize(
    ("enabled", "expected"),
    [
        (True, ["До", "Город: Минск", "После"]),
        (False, ["До", "После"]),
    ],
)
def test_renderer_renders_body_condition_with_markers_split_across_runs(
    enabled, expected
):
    document = Document()
    document.add_paragraph("До")
    opening = document.add_paragraph()
    opening.add_run("{{#if seller.is_")
    opening.add_run("individual_entrepreneur}}")
    document.add_paragraph("Город: {{ seller.city }}")
    closing = document.add_paragraph()
    closing.add_run("{{/if seller.is_individual_")
    closing.add_run("entrepreneur}}")
    document.add_paragraph("После")

    rendered = NativeDocxRenderer().render(
        _conditional_template(_save(document)),
        RenderContext(
            values={"seller.city": "Минск"},
            conditions={"seller.is_individual_entrepreneur": enabled},
        ),
    )

    result = Document(BytesIO(rendered.content))
    assert [paragraph.text for paragraph in result.paragraphs] == expected


@pytest.mark.parametrize(
    ("enabled", "expected_rows"),
    [
        (True, [["Всегда", ""], ["Филиал", "Минск"]]),
        (False, [["Всегда", ""]]),
    ],
)
def test_renderer_renders_whole_table_row_conditions(enabled, expected_rows):
    document = Document()
    table = document.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Всегда"
    opening = table.rows[1].cells[0].paragraphs[0]
    opening.add_run("{{#if seller.is_individual_")
    opening.add_run("entrepreneur}}")
    table.rows[2].cells[0].text = "Филиал"
    table.rows[2].cells[1].text = "{{ seller.city }}"
    table.rows[3].cells[0].text = "{{/if seller.is_individual_entrepreneur}}"

    rendered = NativeDocxRenderer().render(
        _conditional_template(_save(document)),
        RenderContext(
            values={"seller.city": "Минск"},
            conditions={"seller.is_individual_entrepreneur": enabled},
        ),
    )

    result = Document(BytesIO(rendered.content))
    assert [[cell.text for cell in row.cells] for row in result.tables[0].rows] == (
        expected_rows
    )


def test_renderer_supports_nested_conditions_in_headers_and_footers():
    outer = "seller.is_individual_entrepreneur"
    inner = "seller.has_branch"
    document = Document()
    header = document.sections[0].header
    header.add_paragraph(f"{{{{#if {outer}}}}}")
    header.add_paragraph("Реквизиты ИП")
    header.add_paragraph(f"{{{{#if {inner}}}}}")
    header.add_paragraph("Филиал: {{ seller.city }}")
    header.add_paragraph(f"{{{{/if {inner}}}}}")
    header.add_paragraph(f"{{{{/if {outer}}}}}")
    footer = document.sections[0].footer
    footer.add_paragraph(f"{{{{#if {inner}}}}}")
    footer.add_paragraph("Скрытый филиал")
    footer.add_paragraph(f"{{{{/if {inner}}}}}")

    rendered = NativeDocxRenderer().render(
        _conditional_template(_save(document), conditions=frozenset({outer, inner})),
        RenderContext(
            values={"seller.city": "Минск"},
            conditions={outer: True, inner: False},
        ),
    )

    result = Document(BytesIO(rendered.content))
    assert [paragraph.text for paragraph in result.sections[0].header.paragraphs] == [
        "",
        "Реквизиты ИП",
    ]
    assert [paragraph.text for paragraph in result.sections[0].footer.paragraphs] == [
        ""
    ]


def test_renderer_handles_first_and_even_page_header_footer_stories():
    condition = "seller.is_individual_entrepreneur"
    document = Document()
    document.sections[0].different_first_page_header_footer = True
    document.settings.odd_and_even_pages_header_footer = True
    for name in (
        "first_page_header",
        "even_page_header",
        "first_page_footer",
        "even_page_footer",
    ):
        area = getattr(document.sections[0], name)
        area.add_paragraph("Город: {{ seller.city }}")
        area.add_paragraph(f"{{{{#if {condition}}}}}")
        area.add_paragraph("Скрытая формулировка")
        area.add_paragraph(f"{{{{/if {condition}}}}}")
    source = _save(document)
    renderer = NativeDocxRenderer()

    assert renderer.discover_placeholders(source) == {"seller.city"}
    assert renderer.discover_conditions(source) == {condition}

    rendered = renderer.render(
        _conditional_template(source),
        RenderContext(
            values={"seller.city": "Минск"},
            conditions={condition: False},
        ),
    )
    result = Document(BytesIO(rendered.content))
    for name in (
        "first_page_header",
        "even_page_header",
        "first_page_footer",
        "even_page_footer",
    ):
        text = "\n".join(
            paragraph.text for paragraph in getattr(result.sections[0], name).paragraphs
        )
        assert "Город: Минск" in text
        assert "Скрытая формулировка" not in text
        assert "{{#if" not in text


def test_renderer_discovers_conditions_separately_from_scalar_placeholders():
    document = Document()
    document.add_paragraph("{{#if seller.is_individual_entrepreneur}}")
    document.add_paragraph("{{ seller.city }}")
    document.add_paragraph("{{/if seller.is_individual_entrepreneur}}")
    source = _save(document)
    renderer = NativeDocxRenderer()

    assert renderer.discover_conditions(source) == {"seller.is_individual_entrepreneur"}
    assert renderer.discover_placeholders(source) == {"seller.city"}


@pytest.mark.parametrize(
    ("paragraphs", "issue_code"),
    [
        (
            ["{{#if seller.is_individual_entrepreneur}}", "Текст"],
            "unbalanced_condition_marker",
        ),
        (
            [
                "{{#if seller.is_individual_entrepreneur}}",
                "{{/if seller.has_branch}}",
            ],
            "mismatched_condition_marker",
        ),
        (
            [
                "{{#if seller.kind == 'ip'}}",
                "{{/if seller.is_individual_entrepreneur}}",
            ],
            "malformed_condition_marker",
        ),
        (
            [
                "Префикс {{#if seller.is_individual_entrepreneur}}",
                "{{/if seller.is_individual_entrepreneur}}",
            ],
            "condition_marker_placement",
        ),
    ],
)
def test_renderer_fails_closed_for_invalid_condition_markers(paragraphs, issue_code):
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    template = _conditional_template(_save(document))
    context = RenderContext(
        values={}, conditions={"seller.is_individual_entrepreneur": True}
    )

    with pytest.raises(TemplateValidationError) as error:
        NativeDocxRenderer().render(template, context)

    assert issue_code in {issue.code for issue in error.value.result.issues}


def test_renderer_rejects_unknown_condition_and_context_condition_leakage():
    document = Document()
    document.add_paragraph("{{#if seller.has_branch}}")
    document.add_paragraph("Филиал")
    document.add_paragraph("{{/if seller.has_branch}}")

    with pytest.raises(TemplateValidationError) as unknown_error:
        NativeDocxRenderer().render(
            _conditional_template(_save(document)),
            RenderContext(values={}, conditions={"seller.has_branch": True}),
        )
    assert "unknown_condition" in {
        issue.code for issue in unknown_error.value.result.issues
    }

    approved_document = Document()
    approved_document.add_paragraph("{{#if seller.is_individual_entrepreneur}}")
    approved_document.add_paragraph("{{/if seller.is_individual_entrepreneur}}")
    with pytest.raises(ContextFieldError, match="internal.secret"):
        NativeDocxRenderer().render(
            _conditional_template(_save(approved_document)),
            RenderContext(
                values={},
                conditions={
                    "seller.is_individual_entrepreneur": True,
                    "internal.secret": False,
                },
            ),
        )


def test_condition_context_is_frozen_and_accepts_only_booleans():
    context = RenderContext(
        values={}, conditions={"seller.is_individual_entrepreneur": True}
    )

    with pytest.raises(TypeError):
        context.conditions["seller.is_individual_entrepreneur"] = False
    with pytest.raises(TypeError, match="must be a bool"):
        RenderContext(values={}, conditions={"seller.is_individual_entrepreneur": 1})


def test_table_condition_marker_cannot_share_its_row_with_other_content():
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "{{#if seller.is_individual_entrepreneur}}"
    table.rows[0].cells[1].text = "Лишний текст"
    table.rows[1].cells[0].text = "{{/if seller.is_individual_entrepreneur}}"

    with pytest.raises(TemplateValidationError) as error:
        NativeDocxRenderer().render(
            _conditional_template(_save(document)),
            RenderContext(
                values={},
                conditions={"seller.is_individual_entrepreneur": True},
            ),
        )

    assert "condition_marker_placement" in {
        issue.code for issue in error.value.result.issues
    }


def test_unconfigured_pdf_converter_is_explicitly_unavailable():
    converter = UnavailablePdfConverter()

    assert converter.health().available is False
    with pytest.raises(
        PdfConversionUnavailableError, match="No PDF conversion provider"
    ):
        converter.convert_docx(b"not-a-real-docx")


def test_gotenberg_health_reports_an_unreachable_explicit_endpoint(monkeypatch):
    import httpx

    def _unreachable(*args, **kwargs):
        raise httpx.ConnectError("connection refused")

    monkeypatch.setattr(
        "modules.documents.infrastructure.renderers.pdf.httpx.get", _unreachable
    )

    health = GotenbergPdfConverter("http://gotenberg.example.invalid").health()

    assert health.available is False
    assert health.provider == "gotenberg"
    assert "unreachable" in health.detail


def test_gotenberg_rejects_success_response_that_is_not_a_pdf(monkeypatch):
    import httpx

    request = httpx.Request("POST", "http://gotenberg.example.invalid")
    monkeypatch.setattr(
        "modules.documents.infrastructure.renderers.pdf.httpx.post",
        lambda *args, **kwargs: httpx.Response(
            200, content=b"<html>proxy error</html>", request=request
        ),
    )

    with pytest.raises(PdfConversionError, match="not a PDF"):
        GotenbergPdfConverter("http://gotenberg.example.invalid").convert_docx(
            b"PK-docx",
            filename="contract.docx",
        )
