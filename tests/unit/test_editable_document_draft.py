from io import BytesIO

import pytest
from docx import Document

from models import DocumentTemplate, DocumentTemplateVersion
from modules.documents.application.editable_draft import (
    EditableDraftError,
    finalize_editable_draft,
    render_editable_draft,
    official_placeholder_counts,
    validate_editable_draft,
)


def _docx(*paragraphs: str) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _text(content: bytes) -> str:
    document = Document(BytesIO(content))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_editable_draft_renders_business_data_but_keeps_official_markers():
    source = _docx(
        "Договор № {{ document.official_full_number }} от {{ document.issued_on }}",
        "Покупатель: {{ customer.full_name }}",
    )
    template = DocumentTemplate(id=7, name="Договор", doc_type="contract")
    version = DocumentTemplateVersion(
        id=9,
        template_id=7,
        version=1,
        renderer="docx",
        source_storage_key="test",
        checksum_sha256="0" * 64,
        source_filename="contract.docx",
        placeholder_schema={
            "fields": [
                "document.official_full_number",
                "document.issued_on",
                "customer.full_name",
            ],
            "conditions": [],
            "tables": [],
        },
    )

    editable = render_editable_draft(
        template=template,
        version=version,
        source=source,
        snapshot={
            "values": {
                "document.official_full_number": "",
                "document.issued_on": "02.09.2026",
                "customer.full_name": "ООО Клиент",
            },
            "conditions": {},
            "table_rows": {},
        },
    )

    text = _text(editable)
    assert "ООО Клиент" in text
    assert "{{ document.official_full_number }}" in text
    assert "{{ document.issued_on }}" in text

    final = finalize_editable_draft(
        source=editable,
        placeholder_schema=version.placeholder_schema,
        official_values={
            "document.official_full_number": "Д-2026-056",
            "document.issued_on": "02.09.2026",
        },
    )
    final_text = _text(final)
    assert "Д-2026-056" in final_text
    assert "02.09.2026" in final_text
    assert "{{" not in final_text


def test_editable_draft_rejects_deleted_official_marker_before_issue():
    source = _docx("Договор без служебного номера")

    with pytest.raises(EditableDraftError, match="удалены служебные поля"):
        validate_editable_draft(
            source=source,
            placeholder_schema={"fields": ["document.official_full_number"]},
        )


def test_editable_draft_rejects_new_unknown_placeholder():
    source = _docx(
        "Договор № {{ document.official_full_number }}",
        "{{ customer.bank_account }}",
    )

    with pytest.raises(EditableDraftError, match="неизвестные плейсхолдеры"):
        validate_editable_draft(
            source=source,
            placeholder_schema={"fields": ["document.official_full_number"]},
        )


def test_editable_draft_rejects_deleted_duplicate_official_marker():
    original = _docx(
        "Договор № {{ document.official_full_number }}",
        "В приложении к договору № {{ document.official_full_number }}",
    )
    edited = _docx("Договор № {{ document.official_full_number }}")
    schema = {"fields": ["document.official_full_number"]}

    required_counts = official_placeholder_counts(
        source=original,
        placeholder_schema=schema,
    )

    with pytest.raises(EditableDraftError, match="изменено количество"):
        validate_editable_draft(
            source=edited,
            placeholder_schema=schema,
            required_placeholder_counts=required_counts,
        )
