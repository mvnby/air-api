"""Restricted conditional sections for native DOCX templates."""

from __future__ import annotations

from dataclasses import dataclass
import re
from typing import Iterable, Mapping

from docx.oxml.ns import qn
from docx.table import _Cell, _Row, Table

from .contracts import TemplateValidationIssue


_IDENTIFIER = r"[a-z][a-z0-9_]*(?:\.[a-z][a-z0-9_]*)*"
_CONDITION_PATTERN = re.compile(rf"{{{{\s*(#if|/if)\s+({_IDENTIFIER})\s*}}}}")
_RAW_PLACEHOLDER_PATTERN = re.compile(r"{{(.*?)}}", re.DOTALL)


def iter_section_story_areas(section):
    """Yield every Word header/footer story exposed by python-docx."""

    for name in (
        "header",
        "first_page_header",
        "even_page_header",
        "footer",
        "first_page_footer",
        "even_page_footer",
    ):
        yield name, getattr(section, name)


@dataclass(frozen=True, slots=True)
class _ConditionMarker:
    kind: str
    name: str
    location: str = ""


class DocxConditionProcessor:
    """Validate and render catalogued boolean DOCX condition markers."""

    def discover(self, document) -> frozenset[str]:
        discovered: set[str] = set()
        for paragraphs in self._all_paragraph_collections(document):
            for paragraph in paragraphs:
                markers, _malformed = self._parse_markers(self._text(paragraph._p))
                discovered.update(marker.name for marker in markers)
        return frozenset(discovered)

    def validate(
        self, document, condition_catalog: frozenset[str]
    ) -> tuple[list[TemplateValidationIssue], set[str]]:
        issues: list[TemplateValidationIssue] = []
        used: set[str] = set()
        for scope, candidates in self._iter_candidates(document):
            stack: list[_ConditionMarker] = []
            for texts, exact, location, placement_message in candidates:
                parsed: list[_ConditionMarker] = []
                for text in texts:
                    found, malformed = self._parse_markers(text)
                    parsed.extend(found)
                    issues.extend(
                        self._issue(
                            "malformed_condition_marker",
                            f"Malformed condition marker {raw!r}",
                            location,
                            raw,
                        )
                        for raw in malformed
                    )
                if not parsed:
                    continue
                if exact is None or len(parsed) != 1:
                    issues.append(
                        self._issue(
                            "condition_marker_placement", placement_message, location
                        )
                    )
                    continue
                marker = _ConditionMarker(exact.kind, exact.name, location)
                used.add(marker.name)
                if marker.name not in condition_catalog:
                    issues.append(
                        self._issue(
                            "unknown_condition",
                            f"Condition '{marker.name}' is not in this template version's catalogue",
                            location,
                            marker.name,
                        )
                    )
                self._advance_stack(stack, marker, issues)
            issues.extend(
                self._issue(
                    "unbalanced_condition_marker",
                    f"Condition '{marker.name}' has no matching closing marker",
                    marker.location or scope,
                    marker.name,
                )
                for marker in stack
            )
        return issues, used

    def render(self, document, conditions: Mapping[str, bool]) -> None:
        for _paragraphs, _location, container in self._paragraph_scopes(document):
            active: list[bool] = []
            for child in list(container):
                marker = (
                    self._exact_marker(self._text(child))
                    if child.tag == qn("w:p")
                    else None
                )
                if marker is not None:
                    container.remove(child)
                    self._update_active(active, marker, conditions)
                elif False in active:
                    container.remove(child)

        for table, _location in list(self._iter_tables(document)):
            active = []
            for row in list(table.rows):
                marker = self._exact_row_marker(row)
                if marker is not None:
                    table._tbl.remove(row._tr)
                    self._update_active(active, marker, conditions)
                elif False in active:
                    table._tbl.remove(row._tr)

    @staticmethod
    def _update_active(
        active: list[bool], marker: _ConditionMarker, conditions: Mapping[str, bool]
    ) -> None:
        if marker.kind == "#if":
            active.append(conditions[marker.name])
        else:
            active.pop()

    @staticmethod
    def _issue(
        code: str, message: str, location: str, placeholder: str | None = None
    ) -> TemplateValidationIssue:
        return TemplateValidationIssue(code, message, location, placeholder)

    @classmethod
    def _advance_stack(
        cls,
        stack: list[_ConditionMarker],
        marker: _ConditionMarker,
        issues: list[TemplateValidationIssue],
    ) -> None:
        if marker.kind == "#if":
            stack.append(marker)
        elif not stack:
            issues.append(
                cls._issue(
                    "unbalanced_condition_marker",
                    f"Closing condition '{marker.name}' has no matching opening marker",
                    marker.location,
                    marker.name,
                )
            )
        else:
            opening = stack.pop()
            if opening.name != marker.name:
                issues.append(
                    cls._issue(
                        "mismatched_condition_marker",
                        f"Condition '{marker.name}' closes '{opening.name}'",
                        marker.location,
                        marker.name,
                    )
                )

    def _iter_candidates(self, document):
        paragraph_message = "Condition marker must be the only content of its paragraph"
        row_message = (
            "Condition marker must be the only content of its entire table row"
        )
        for paragraphs, scope, _container in self._paragraph_scopes(document):
            yield (
                scope,
                [
                    (
                        (self._text(paragraph._p),),
                        self._exact_marker(self._text(paragraph._p)),
                        f"{scope}.paragraph[{index}]",
                        paragraph_message,
                    )
                    for index, paragraph in enumerate(paragraphs)
                ],
            )
        for table, scope in self._iter_tables(document):
            yield (
                scope,
                [
                    (
                        tuple(
                            self._text(p._p) for p in self._direct_row_paragraphs(row)
                        ),
                        self._exact_row_marker(row),
                        f"{scope}.row[{index}]",
                        row_message,
                    )
                    for index, row in enumerate(table.rows)
                ],
            )

    def _all_paragraph_collections(self, document):
        for paragraphs, _scope, _container in self._paragraph_scopes(document):
            yield paragraphs
        for table, _scope in self._iter_tables(document):
            for row in table.rows:
                yield self._direct_row_paragraphs(row)

    @staticmethod
    def _paragraph_scopes(document):
        scopes = [(document.paragraphs, "body", document.element.body)]
        seen: set[object] = {document.element.body}
        for section_index, section in enumerate(document.sections):
            for name, area in iter_section_story_areas(section):
                if area._element in seen:
                    continue
                seen.add(area._element)
                scopes.append(
                    (
                        area.paragraphs,
                        f"section[{section_index}].{name}",
                        area._element,
                    )
                )
        yield from scopes

    @classmethod
    def _parse_markers(cls, text: str) -> tuple[list[_ConditionMarker], list[str]]:
        valid_matches = list(_CONDITION_PATTERN.finditer(text))
        valid_spans = {(match.start(), match.end()) for match in valid_matches}
        malformed = [
            match.group(0)
            for match in _RAW_PLACEHOLDER_PATTERN.finditer(text)
            if match.group(1).strip().startswith(("#if", "/if"))
            and (match.start(), match.end()) not in valid_spans
        ]
        for prefix in ("{{#if", "{{/if"):
            start = text.find(prefix)
            if start >= 0 and "}}" not in text[start:]:
                malformed.append(text[start:])
        return (
            [
                _ConditionMarker(match.group(1), match.group(2))
                for match in valid_matches
            ],
            malformed,
        )

    @staticmethod
    def _exact_marker(text: str) -> _ConditionMarker | None:
        match = _CONDITION_PATTERN.fullmatch(text.strip())
        return (
            _ConditionMarker(match.group(1), match.group(2))
            if match is not None
            else None
        )

    def _exact_row_marker(self, row: _Row) -> _ConditionMarker | None:
        cells = self._unique_cells(row)
        if any(cell.tables for cell in cells):
            return None
        nonempty = [
            text
            for paragraph in self._direct_row_paragraphs(row)
            if (text := self._text(paragraph._p)).strip()
        ]
        return self._exact_marker(nonempty[0]) if len(nonempty) == 1 else None

    @staticmethod
    def _unique_cells(row: _Row) -> list[_Cell]:
        cells: list[_Cell] = []
        seen: set[object] = set()
        for cell in row.cells:
            if cell._tc not in seen:
                seen.add(cell._tc)
                cells.append(cell)
        return cells

    def _direct_row_paragraphs(self, row: _Row):
        paragraphs = []
        seen: set[object] = set()
        for cell in self._unique_cells(row):
            for paragraph in cell.paragraphs:
                if paragraph._p not in seen:
                    seen.add(paragraph._p)
                    paragraphs.append(paragraph)
        return paragraphs

    def _iter_tables(self, document):
        seen: set[object] = set()
        yield from self._iter_table_collection(document.tables, "body", seen)
        for section_index, section in enumerate(document.sections):
            for name, area in iter_section_story_areas(section):
                yield from self._iter_table_collection(
                    area.tables, f"section[{section_index}].{name}", seen
                )

    def _iter_table_collection(
        self, tables: Iterable[Table], location: str, seen: set[object]
    ):
        for table_index, table in enumerate(tables):
            if table._tbl in seen:
                continue
            seen.add(table._tbl)
            table_location = f"{location}.table[{table_index}]"
            yield table, table_location
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(self._unique_cells(row)):
                    yield from self._iter_table_collection(
                        cell.tables,
                        f"{table_location}.row[{row_index}].cell[{cell_index}]",
                        seen,
                    )

    @staticmethod
    def _text(element) -> str:
        return "".join(node.text or "" for node in element.iter(qn("w:t")))
