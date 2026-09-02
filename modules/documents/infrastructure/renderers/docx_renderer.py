"""Safe native DOCX rendering without executing template code."""

from __future__ import annotations

from copy import deepcopy
from collections import Counter
from io import BytesIO
import re
from typing import Iterable, Mapping

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, _Row, Table
from docx.text.paragraph import Paragraph

from .contracts import (
    DocumentTemplateVersion,
    RenderContext,
    TableBlockSpec,
    TemplateValidationError,
    TemplateValidationIssue,
    TemplateValidationResult,
    make_rendered_docx,
)
from .docx_conditions import DocxConditionProcessor, iter_section_story_areas


_PLACEHOLDER_PATTERN = re.compile(r"{{\s*([a-z][a-z0-9_]*(?:\.[a-z][a-z0-9_]*)*)\s*}}")
_RAW_PLACEHOLDER_PATTERN = re.compile(r"{{(.*?)}}", re.DOTALL)


class ContextFieldError(ValueError):
    """The context contains data which the selected template has not approved."""


class NativeDocxRenderer:
    """Render DOCX from a whitelist rather than an expression language.

    Templates can contain only approved identifiers such as
    ``{{ customer.full_name }}``. No functions, filters, attribute traversal,
    or Python/Jinja expressions are evaluated.
    """

    def discover_placeholders(self, source: bytes) -> frozenset[str]:
        """Return syntactically valid placeholder names found in a DOCX.

        Discovery never expands the security catalogue; callers intersect the
        result with their server-owned allowlist before constructing a version.
        """
        return frozenset(self.discover_placeholder_counts(source))

    def discover_placeholder_counts(self, source: bytes) -> dict[str, int]:
        """Count every syntactically valid placeholder occurrence in a DOCX."""
        try:
            document = Document(BytesIO(source))
        except Exception as exc:
            raise ValueError(f"Template is not a readable DOCX: {exc}") from exc
        discovered: Counter[str] = Counter()
        for paragraph, _location, _row_id in self._iter_paragraphs(document):
            parsed, _malformed = self._parse_placeholders(
                self._paragraph_text(paragraph)
            )
            discovered.update(parsed)
        return dict(discovered)

    def discover_conditions(self, source: bytes) -> frozenset[str]:
        """Return safe condition identifiers found in conditional markers."""
        try:
            document = Document(BytesIO(source))
        except Exception as exc:
            raise ValueError(f"Template is not a readable DOCX: {exc}") from exc
        return DocxConditionProcessor().discover(document)

    def validate(
        self,
        template: DocumentTemplateVersion,
        context: RenderContext | None = None,
    ) -> TemplateValidationResult:
        try:
            document = Document(BytesIO(template.source))
        except Exception as exc:  # python-docx exposes several ZIP/XML exceptions
            return TemplateValidationResult(
                (
                    TemplateValidationIssue(
                        "invalid_docx",
                        f"Template is not a readable DOCX: {exc}",
                        "document",
                    ),
                )
            )

        issues: list[TemplateValidationIssue] = []
        blocks = {block.name: block for block in template.table_blocks}
        row_block_by_field = {
            field: block
            for block in template.table_blocks
            for field in block.row_fields
        }
        known = set(template.field_catalog) | set(blocks) | set(row_block_by_field)
        anchor_locations: dict[str, list[str]] = {name: [] for name in blocks}
        row_fields_by_row: dict[object, set[str]] = {}
        anchored_rows: dict[object, set[str]] = {}
        scalar_placeholders: set[str] = set()

        condition_issues, used_conditions = DocxConditionProcessor().validate(
            document, template.condition_catalog
        )
        issues.extend(condition_issues)

        for paragraph, location, row_id in self._iter_paragraphs(document):
            parsed, malformed = self._parse_placeholders(
                self._paragraph_text(paragraph)
            )
            for raw in malformed:
                issues.append(
                    TemplateValidationIssue(
                        "malformed_placeholder",
                        f"Malformed placeholder {raw!r}",
                        location,
                        raw,
                    )
                )
            for placeholder in parsed:
                if placeholder not in known:
                    issues.append(
                        TemplateValidationIssue(
                            "unknown_placeholder",
                            f"Placeholder '{placeholder}' is not in this template version's catalogue",
                            location,
                            placeholder,
                        )
                    )
                elif placeholder in blocks:
                    if row_id is None:
                        issues.append(
                            TemplateValidationIssue(
                                "table_anchor_outside_table",
                                f"Table anchor '{placeholder}' must be inside a DOCX table row",
                                location,
                                placeholder,
                            )
                        )
                    else:
                        anchor_locations[placeholder].append(location)
                        anchored_rows.setdefault(row_id, set()).add(placeholder)
                elif placeholder in row_block_by_field:
                    if row_id is None:
                        issues.append(
                            TemplateValidationIssue(
                                "row_placeholder_outside_table",
                                f"Row placeholder '{placeholder}' must be inside its table anchor row",
                                location,
                                placeholder,
                            )
                        )
                    else:
                        row_fields_by_row.setdefault(row_id, set()).add(placeholder)
                else:
                    scalar_placeholders.add(placeholder)

        for name, locations in anchor_locations.items():
            if not locations:
                issues.append(
                    TemplateValidationIssue(
                        "missing_table_anchor",
                        f"Table block '{name}' has no '{{{{ {name} }}}}' anchor in the template",
                        "document",
                        name,
                    )
                )
            elif len(locations) > 1:
                issues.append(
                    TemplateValidationIssue(
                        "duplicate_table_anchor",
                        f"Table block '{name}' has {len(locations)} anchors; exactly one is required",
                        ", ".join(locations),
                        name,
                    )
                )
        for row_id, fields in row_fields_by_row.items():
            anchors = anchored_rows.get(row_id, set())
            for field_name in fields:
                expected_anchor = row_block_by_field[field_name].name
                if expected_anchor not in anchors:
                    issues.append(
                        TemplateValidationIssue(
                            "row_placeholder_without_anchor",
                            f"Row placeholder '{field_name}' must share a row with '{{{{ {expected_anchor} }}}}'",
                            f"table-row:{row_id}",
                            field_name,
                        )
                    )
        for row_id, anchors in anchored_rows.items():
            row_fields = row_fields_by_row.get(row_id, set())
            for anchor in anchors:
                if any(
                    row_block_by_field[field].name != anchor for field in row_fields
                ):
                    issues.append(
                        TemplateValidationIssue(
                            "mixed_table_blocks",
                            f"Anchor '{anchor}' cannot share a row with fields from another table block",
                            f"table-row:{row_id}",
                            anchor,
                        )
                    )

        if context is not None:
            issues.extend(
                self._validate_context(
                    template,
                    context,
                    scalar_placeholders,
                    row_fields_by_row,
                    used_conditions,
                )
            )
        return TemplateValidationResult(tuple(issues))

    def render(self, template: DocumentTemplateVersion, context: RenderContext):
        validation = self.validate(template, context)
        if not validation.is_valid:
            if all(issue.code.startswith("context_") for issue in validation.issues):
                raise ContextFieldError(
                    "; ".join(issue.message for issue in validation.issues)
                )
            raise TemplateValidationError(validation)

        document = Document(BytesIO(template.source))
        DocxConditionProcessor().render(document, context.conditions)
        for table, _ in self._iter_tables(document):
            self._render_table_blocks(table, template.table_blocks, context)
        for paragraph, _, _ in self._iter_paragraphs(document):
            self._replace_placeholders(paragraph, context.values)

        self._remove_empty_trailing_paragraph_after_table(document)

        output = BytesIO()
        document.save(output)
        return make_rendered_docx(output.getvalue(), template)

    @staticmethod
    def _remove_empty_trailing_paragraph_after_table(document) -> None:
        """Drop a non-content paragraph that would create a blank final page.

        Word requires a paragraph after a table while editing, and Google Docs
        exports one at the end of the document.  When the signature table fills
        the previous page, that otherwise invisible paragraph can spill onto a
        blank page.  Keep paragraphs containing any renderable OOXML content;
        only the final, truly empty paragraph immediately after a table is safe
        to remove from the generated artifact.
        """
        body = document._element.body
        content = [child for child in body.iterchildren() if child.tag != qn("w:sectPr")]
        if len(content) < 2:
            return

        previous, trailing = content[-2:]
        if previous.tag != qn("w:tbl") or trailing.tag != qn("w:p"):
            return

        rendered_tags = {
            qn("w:t"),
            qn("w:tab"),
            qn("w:br"),
            qn("w:drawing"),
            qn("w:object"),
            qn("w:pict"),
            qn("w:sym"),
            qn("w:fldChar"),
            qn("w:instrText"),
            qn("w:noBreakHyphen"),
        }
        if any(node.tag in rendered_tags for node in trailing.iter()):
            return

        body.remove(trailing)

    def _validate_context(
        self,
        template: DocumentTemplateVersion,
        context: RenderContext,
        scalar_placeholders: set[str],
        row_fields_by_row: Mapping[object, set[str]],
        used_conditions: set[str],
    ) -> list[TemplateValidationIssue]:
        issues: list[TemplateValidationIssue] = []
        for field_name in sorted(set(context.values) - set(template.field_catalog)):
            issues.append(
                TemplateValidationIssue(
                    "context_unknown_field",
                    f"Context field '{field_name}' is not approved by this template version",
                    "context.values",
                    field_name,
                )
            )
        for field_name in sorted(scalar_placeholders - set(context.values)):
            issues.append(
                TemplateValidationIssue(
                    "context_missing_field",
                    f"Context does not provide required field '{field_name}'",
                    "context.values",
                    field_name,
                )
            )

        for condition_name in sorted(
            set(context.conditions) - set(template.condition_catalog)
        ):
            issues.append(
                TemplateValidationIssue(
                    "context_unknown_condition",
                    f"Context condition '{condition_name}' is not approved by this template version",
                    "context.conditions",
                    condition_name,
                )
            )
        for condition_name in sorted(used_conditions - set(context.conditions)):
            issues.append(
                TemplateValidationIssue(
                    "context_missing_condition",
                    f"Context does not provide required condition '{condition_name}'",
                    "context.conditions",
                    condition_name,
                )
            )

        blocks = {block.name: block for block in template.table_blocks}
        for table_name in sorted(set(context.table_rows) - set(blocks)):
            issues.append(
                TemplateValidationIssue(
                    "context_unknown_table",
                    f"Context table '{table_name}' is not approved by this template version",
                    "context.table_rows",
                    table_name,
                )
            )
        used_fields: dict[str, set[str]] = {name: set() for name in blocks}
        for fields in row_fields_by_row.values():
            for field_name in fields:
                for block in template.table_blocks:
                    if field_name in block.row_fields:
                        used_fields[block.name].add(field_name)
                        break
        for table_name, rows in context.table_rows.items():
            block = blocks.get(table_name)
            if block is None:
                continue
            for index, row in enumerate(rows):
                for field_name in sorted(set(row) - set(block.row_fields)):
                    issues.append(
                        TemplateValidationIssue(
                            "context_unknown_row_field",
                            f"Table '{table_name}' row {index + 1} has unapproved field '{field_name}'",
                            f"context.table_rows.{table_name}[{index}]",
                            field_name,
                        )
                    )
                for field_name in sorted(used_fields[table_name] - set(row)):
                    issues.append(
                        TemplateValidationIssue(
                            "context_missing_row_field",
                            f"Table '{table_name}' row {index + 1} lacks required field '{field_name}'",
                            f"context.table_rows.{table_name}[{index}]",
                            field_name,
                        )
                    )
        return issues

    def _render_table_blocks(
        self,
        table: Table,
        blocks: Iterable[TableBlockSpec],
        context: RenderContext,
    ) -> None:
        for block in blocks:
            template_rows = [
                row
                for row in table.rows
                if block.name in self._placeholders_in_row(row)
            ]
            if not template_rows:
                continue
            template_row = template_rows[0]
            insert_at = list(table._tbl).index(template_row._tr)
            for row_values in context.table_rows.get(block.name, ()):
                cloned_tr = deepcopy(template_row._tr)
                table._tbl.insert(insert_at, cloned_tr)
                insert_at += 1
                fields = dict(context.values)
                fields.update(row_values)
                fields[block.name] = ""
                for cell in _Row(cloned_tr, table).cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders(paragraph, fields)
            table._tbl.remove(template_row._tr)

    def _placeholders_in_row(self, row: _Row) -> set[str]:
        return {
            field
            for cell in row.cells
            for paragraph in cell.paragraphs
            for field in self._parse_placeholders(self._paragraph_text(paragraph))[0]
        }

    @staticmethod
    def _paragraph_text(paragraph: Paragraph) -> str:
        return "".join(node.text or "" for node in paragraph._p.iter(qn("w:t")))

    @staticmethod
    def _parse_placeholders(text: str) -> tuple[list[str], list[str]]:
        valid_matches = list(_PLACEHOLDER_PATTERN.finditer(text))
        valid_spans = {(match.start(), match.end()) for match in valid_matches}
        malformed = [
            match.group(0)
            for match in _RAW_PLACEHOLDER_PATTERN.finditer(text)
            if (match.start(), match.end()) not in valid_spans
            and not match.group(1).strip().startswith(("#if", "/if"))
        ]
        if "{{" in text and not _RAW_PLACEHOLDER_PATTERN.search(text):
            malformed.append("{{")
        return [match.group(1) for match in valid_matches], malformed

    def _replace_placeholders(
        self, paragraph: Paragraph, values: Mapping[str, str]
    ) -> None:
        nodes = list(paragraph._p.iter(qn("w:t")))
        source = "".join(node.text or "" for node in nodes)
        matches = list(_PLACEHOLDER_PATTERN.finditer(source))
        if not matches:
            return
        positions: list[tuple[int, int]] = []
        cursor = 0
        for node in nodes:
            end = cursor + len(node.text or "")
            positions.append((cursor, end))
            cursor = end
        for match in reversed(matches):
            field = match.group(1)
            if field not in values:
                continue
            start_node, start_offset = self._locate(positions, match.start())
            end_node, end_offset = self._locate(positions, match.end(), is_end=True)
            replacement = values[field]
            if start_node == end_node:
                text = nodes[start_node].text or ""
                self._set_text(
                    nodes[start_node],
                    text[:start_offset] + replacement + text[end_offset:],
                )
            else:
                start_text = nodes[start_node].text or ""
                end_text = nodes[end_node].text or ""
                self._set_text(
                    nodes[start_node], start_text[:start_offset] + replacement
                )
                for index in range(start_node + 1, end_node):
                    self._set_text(nodes[index], "")
                self._set_text(nodes[end_node], end_text[end_offset:])

    @staticmethod
    def _set_text(node, value: str) -> None:
        parts = value.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        NativeDocxRenderer._set_text_node(node, parts[0])
        if len(parts) == 1:
            return

        parent = node.getparent()
        position = parent.index(node) + 1
        for part in parts[1:]:
            parent.insert(position, OxmlElement("w:br"))
            position += 1
            text_node = OxmlElement("w:t")
            NativeDocxRenderer._set_text_node(text_node, part)
            parent.insert(position, text_node)
            position += 1

    @staticmethod
    def _set_text_node(node, value: str) -> None:
        node.text = value
        space_attr = qn("xml:space")
        if value.startswith(" ") or value.endswith(" "):
            node.set(space_attr, "preserve")
        elif space_attr in node.attrib:
            del node.attrib[space_attr]

    @staticmethod
    def _locate(
        positions: list[tuple[int, int]], position: int, *, is_end: bool = False
    ) -> tuple[int, int]:
        for index, (start, end) in enumerate(positions):
            if start <= position < end or (is_end and position == end):
                return index, position - start
        index = len(positions) - 1
        return index, positions[index][1] - positions[index][0]

    def _iter_paragraphs(self, document):
        # Keep the XML nodes themselves alive while walking. Python may reuse
        # ``id()`` values for short-lived wrapper objects, which previously made
        # valid table anchors disappear nondeterministically during validation.
        seen: set[object] = set()
        for paragraph in document.paragraphs:
            if paragraph._p not in seen:
                seen.add(paragraph._p)
                yield paragraph, "body", None
        for table, location in self._iter_tables(document):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    yield from self._iter_cell_paragraphs(
                        cell,
                        f"{location}.row[{row_index}].cell[{cell_index}]",
                        row._tr,
                        seen,
                    )
        for section_index, section in enumerate(document.sections):
            for name, area in iter_section_story_areas(section):
                for paragraph in area.paragraphs:
                    if paragraph._p not in seen:
                        seen.add(paragraph._p)
                        yield paragraph, f"section[{section_index}].{name}", None
                for table_index, table in enumerate(area.tables):
                    for row_index, row in enumerate(table.rows):
                        for cell_index, cell in enumerate(row.cells):
                            yield from self._iter_cell_paragraphs(
                                cell,
                                f"section[{section_index}].{name}.table[{table_index}].row[{row_index}].cell[{cell_index}]",
                                row._tr,
                                seen,
                            )

    def _iter_cell_paragraphs(
        self, cell: _Cell, location: str, row_id: object, seen: set[object]
    ):
        for paragraph in cell.paragraphs:
            if paragraph._p not in seen:
                seen.add(paragraph._p)
                yield paragraph, location, row_id
        for table_index, table in enumerate(cell.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, nested_cell in enumerate(row.cells):
                    yield from self._iter_cell_paragraphs(
                        nested_cell,
                        f"{location}.table[{table_index}].row[{row_index}].cell[{cell_index}]",
                        row._tr,
                        seen,
                    )

    def _iter_tables(self, document):
        seen: set[object] = set()
        yield from self._iter_table_collection(document.tables, "body", seen)
        for section_index, section in enumerate(document.sections):
            for name, area in iter_section_story_areas(section):
                yield from self._iter_table_collection(
                    area.tables,
                    f"section[{section_index}].{name}",
                    seen,
                )

    def _iter_table_collection(
        self, tables: Iterable[Table], location: str, seen: set[object]
    ):
        for table_index, table in enumerate(tables):
            if table._tbl in seen:
                continue
            seen.add(table._tbl)
            table_location = f"{location}.table[{table_index}]"
            yield table, table_location
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    yield from self._iter_table_collection(
                        cell.tables,
                        f"{table_location}.row[{row_index}].cell[{cell_index}]",
                        seen,
                    )
